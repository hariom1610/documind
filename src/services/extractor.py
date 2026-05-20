import io
import logging
from typing import Tuple

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract text from PDF using PyMuPDF (fitz).
    Falls back to OCR if text extraction yields insufficient content (scanned PDFs).
    """
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages_text = []

        for page in doc:
            text = page.get_text("text")
            pages_text.append(text)

        full_text = "\n".join(pages_text).strip()
        doc.close()

        # If extracted text is too sparse, likely a scanned PDF — try OCR
        if len(full_text) < 50:
            logger.info("PDF has minimal text, attempting OCR fallback...")
            full_text = _ocr_pdf_pages(file_bytes)

        return full_text

    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        raise ValueError(f"Failed to extract text from PDF: {str(e)}")


def _ocr_pdf_pages(file_bytes: bytes) -> str:
    """OCR fallback for scanned PDFs — renders each page as image and runs Tesseract in parallel."""
    try:
        import fitz
        import pytesseract
        from PIL import Image
        from concurrent.futures import ThreadPoolExecutor
        import os

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages_img_bytes = []

        # Rendering pages is extremely fast in PyMuPDF (fitz), so we render sequentially
        for page in doc:
            # Render page at 2x zoom for better OCR accuracy
            mat = fitz.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=mat)
            pages_img_bytes.append(pix.tobytes("png"))

        doc.close()

        def ocr_single_page(img_bytes: bytes) -> str:
            try:
                img = Image.open(io.BytesIO(img_bytes))
                return pytesseract.image_to_string(img, config="--oem 3 --psm 6")
            except Exception as page_err:
                logger.error(f"Page OCR processing failed: {page_err}")
                return ""

        # Run Tesseract OCR in parallel across CPU cores. Limit max workers to 4.
        max_workers = min(4, os.cpu_count() or 2)
        logger.info(f"Running parallel PDF OCR using {max_workers} threads for {len(pages_img_bytes)} pages.")
        
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            ocr_texts = list(executor.map(ocr_single_page, pages_img_bytes))

        return "\n".join(t.strip() for t in ocr_texts if t.strip()).strip()

    except Exception as e:
        logger.error(f"PDF OCR fallback failed: {e}")
        return ""



def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract text from DOCX using python-docx.
    Extracts paragraphs, tables, and headers for complete coverage.
    """
    try:
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))
        content_parts = []

        # Extract all paragraphs (includes headings, normal text, lists)
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                content_parts.append(text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    content_parts.append(" | ".join(row_texts))
        for section in doc.sections:
            for container in [section.header, section.footer]:
                if container:
                    for para in container.paragraphs:
                        text = para.text.strip()
                        if text:
                            content_parts.append(text)

        full_text = "\n".join(content_parts).strip()

        if not full_text:
            raise ValueError("No readable text found in DOCX file")

        return full_text

    except ValueError:
        raise
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        raise ValueError(f"Failed to extract text from DOCX: {str(e)}")


def extract_text_from_image(file_bytes: bytes) -> str:
    """
    Extract text from image using Tesseract OCR via pytesseract.
    Applies preprocessing for better accuracy.
    """
    try:
        import pytesseract
        from PIL import Image, ImageEnhance, ImageFilter

        img = Image.open(io.BytesIO(file_bytes))

        # Convert to RGB if needed (handles RGBA, palette mode, etc.)
        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")

        # Preprocessing: sharpen and enhance contrast for better OCR
        img = img.filter(ImageFilter.SHARPEN)
        enhancer = ImageEnhance.Contrast(img)
        img = enhancer.enhance(1.5)

        # Use LSTM OCR engine with automatic page segmentation
        custom_config = "--oem 3 --psm 6"
        text = pytesseract.image_to_string(img, config=custom_config)
        text = text.strip()

        if not text:
            # Try with different page segmentation mode for single-block text
            text = pytesseract.image_to_string(img, config="--oem 3 --psm 3")
            text = text.strip()

        if not text:
            raise ValueError("No text could be extracted from the image")

        return text

    except ValueError:
        raise
    except Exception as e:
        logger.error(f"Image OCR failed: {e}")
        raise ValueError(f"Failed to extract text from image: {str(e)}")


def extract_text(file_bytes: bytes, file_type: str) -> str:
    """
    Main dispatcher: routes to the correct extractor based on file_type.
    Returns extracted text string.
    """
    file_type = file_type.lower().strip()

    extractors = {
        "pdf": extract_text_from_pdf,
        "docx": extract_text_from_docx,
        "image": extract_text_from_image,
    }

    if file_type not in extractors:
        raise ValueError(f"Unsupported file type: '{file_type}'. Must be pdf, docx, or image.")

    logger.info(f"Extracting text from {file_type} file ({len(file_bytes)} bytes)")
    text = extractors[file_type](file_bytes)
    logger.info(f"Extracted {len(text)} characters from {file_type}")

    return text
